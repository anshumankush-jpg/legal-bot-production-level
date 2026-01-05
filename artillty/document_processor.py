"""
Document Processor for PLAZA-AI
Handles PDF, DOCX, images, and text extraction with intelligent chunking
"""

import os
import re
import io
import tempfile
from pathlib import Path
from typing import List, Dict, Optional, Any, Union, Tuple
import logging

# Document processing libraries
try:
    import PyPDF2 as pypdf2
except ImportError:
    pypdf2 = None

try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import openpyxl
    OPENPYXL_AVAILABLE = True
except ImportError:
    OPENPYXL_AVAILABLE = False

try:
    import pytesseract
    from PIL import Image
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False

try:
    import easyocr
    EASYOCR_AVAILABLE = True
except ImportError:
    EASYOCR_AVAILABLE = False

# Simple text splitter (replacing LangChain dependency)
class SimpleCharacterTextSplitter:
    """Simple character-based text splitter."""

    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200, separators: List[str] = None):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.separators = separators or ["\n\n", "\n", ". ", " ", ""]

    def split_text(self, text: str) -> List[str]:
        """Split text into chunks."""
        if not text:
            return []

        chunks = []
        start = 0

        while start < len(text):
            # Find the end of this chunk
            end = start + self.chunk_size

            # If we're not at the end, try to find a good break point
            if end < len(text):
                # Look for separators in reverse order of preference
                for separator in self.separators:
                    # Look for separator within the last part of the chunk
                    search_start = max(start, end - 100)  # Don't go too far back
                    sep_index = text.rfind(separator, search_start, end)
                    if sep_index != -1:
                        end = sep_index + len(separator)
                        break

            # Extract chunk
            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)

            # Move start position with overlap
            start = end - self.chunk_overlap

            # Prevent infinite loops
            if start >= len(text):
                break
            if start <= start - self.chunk_overlap:  # No progress
                start = end

        return chunks

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """
    Document processor for PLAZA-AI.

    Features:
    - Multi-format document processing (PDF, DOCX, XLSX, images, text)
    - OCR for scanned documents
    - Intelligent text chunking with overlap
    - Offence number detection and validation
    - Table extraction and processing
    """

    def __init__(
        self,
        chunk_size: int = 1000,
        chunk_overlap: int = 200,
        separators: Optional[List[str]] = None
    ):
        """
        Initialize document processor.

        Args:
            chunk_size: Size of text chunks in characters
            chunk_overlap: Overlap between chunks in characters
            separators: Custom separators for chunking
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.separators = separators or ["\n\n", "\n", ". ", " ", ""]

        # Initialize text splitter
        self.text_splitter = SimpleCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            separators=self.separators
        )

        # Offence number patterns
        self.offence_patterns = [
            r'[Oo]ffence\s*[Nn]o\.?\s*:?\s*(\d{8,12})',
            r'[Oo]ffence\s*[Nn]umber\s*:?\s*(\d{8,12})',
            r'[Tt]icket\s*#?\s*:?\s*(\d{8,12})',
            r'[Cc]itation\s*#?\s*:?\s*(\d{8,12})',
            r'[Vv]iolation\s*#?\s*:?\s*(\d{8,12})',
            r'[Ii]nfraction\s*#?\s*:?\s*(\d{8,12})',
        ]

        logger.info("📄 Document processor initialized")

    def extract_text_from_pdf(self, pdf_path: str) -> str:
        """
        Extract text from PDF using pdfplumber (preferred) or PyPDF2 (fallback).

        Args:
            pdf_path: Path to PDF file

        Returns:
            Extracted text content
        """
        if PDFPLUMBER_AVAILABLE:
            try:
                with pdfplumber.open(pdf_path) as pdf:
                    text_parts = []
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text_parts.append(page_text)
                    return "\n\n".join(text_parts)
            except Exception as e:
                logger.warning(f"pdfplumber extraction failed: {e}")

        # Fallback to PyPDF2
        if pypdf2:
            try:
                with open(pdf_path, 'rb') as file:
                    pdf_reader = pypdf2.PdfReader(file)
                    text_parts = []
                    for page in pdf_reader.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text_parts.append(page_text)
                    return "\n\n".join(text_parts)
            except Exception as e:
                logger.warning(f"PyPDF2 extraction failed: {e}")

        raise ValueError(f"Could not extract text from PDF: {pdf_path}")

    def extract_text_from_docx(self, docx_path: str) -> str:
        """
        Extract text from DOCX document.

        Args:
            docx_path: Path to DOCX file

        Returns:
            Extracted text content
        """
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx not available. Install with: pip install python-docx")

        doc = DocxDocument(docx_path)
        paragraphs = []

        # Extract from paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)

        # Extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    paragraphs.append(" | ".join(row_text))

        return "\n\n".join(paragraphs)

    def extract_text_from_xlsx(self, xlsx_path: str) -> Tuple[str, List[Dict]]:
        """
        Extract text and tables from XLSX file.

        Args:
            xlsx_path: Path to XLSX file

        Returns:
            Tuple of (text_content, tables_list)
        """
        if not OPENPYXL_AVAILABLE:
            raise ImportError("openpyxl not available. Install with: pip install openpyxl")

        workbook = openpyxl.load_workbook(xlsx_path)
        text_parts = []
        tables = []

        for sheet_name in workbook.sheetnames:
            worksheet = workbook[sheet_name]

            # Extract table data
            table_data = []
            for row in worksheet.iter_rows(values_only=True):
                if any(cell for cell in row if cell is not None):
                    table_data.append([str(cell) if cell is not None else "" for cell in row])

            if table_data:
                tables.append({
                    'sheet': sheet_name,
                    'data': table_data
                })

                # Convert table to text
                sheet_text = f"Sheet: {sheet_name}\n"
                for i, row in enumerate(table_data):
                    sheet_text += f"Row {i}: " + " | ".join(row) + "\n"
                text_parts.append(sheet_text)

        return "\n\n".join(text_parts), tables

    def extract_text_from_image(self, image_path: str) -> str:
        """
        Extract text from image using OCR.

        Args:
            image_path: Path to image file

        Returns:
            Extracted text content
        """
        if not OCR_AVAILABLE:
            raise ImportError("OCR not available. Install with: pip install pytesseract Pillow")

        try:
            image = Image.open(image_path)
            # Convert to RGB if necessary
            if image.mode != 'RGB':
                image = image.convert('RGB')

            # Perform OCR
            text = pytesseract.image_to_string(image)
            return text.strip()

        except Exception as e:
            logger.warning(f"OCR failed for {image_path}: {e}")
            return ""

    def extract_text_from_image_easyocr(self, image_path: str) -> str:
        """
        Extract text from image using EasyOCR (alternative to Tesseract).

        Args:
            image_path: Path to image file

        Returns:
            Extracted text content
        """
        if not EASYOCR_AVAILABLE:
            raise ImportError("EasyOCR not available. Install with: pip install easyocr")

        try:
            reader = easyocr.Reader(['en'])
            results = reader.readtext(image_path, detail=0)
            return " ".join(results)

        except Exception as e:
            logger.warning(f"EasyOCR failed for {image_path}: {e}")
            return ""

    def detect_offence_number(self, text: str) -> Optional[str]:
        """
        Detect and validate offence/ticket number from text.

        Args:
            text: Text content to search

        Returns:
            Detected offence number or None
        """
        for pattern in self.offence_patterns:
            matches = re.finditer(pattern, text, re.IGNORECASE)
            for match in matches:
                offence_num = match.group(1).strip()

                # Validate: should be 8-12 digits
                if len(offence_num) >= 8 and len(offence_num) <= 12 and offence_num.isdigit():
                    logger.info(f"✅ Detected offence number: {offence_num}")
                    return offence_num

        return None

    def detect_province(self, text: str) -> Optional[str]:
        """
        Detect Canadian province or US state from text.

        Args:
            text: Text content to search

        Returns:
            Detected province/state or None
        """
        # Canadian provinces
        provinces = {
            'ontario': ['ontario', 'highway traffic act', 'hta'],
            'quebec': ['quebec', 'code de la sécurité routière'],
            'british columbia': ['british columbia', 'motor vehicle act'],
            'alberta': ['alberta', 'traffic safety act'],
            'manitoba': ['manitoba', 'highway traffic act'],
            'saskatchewan': ['saskatchewan', 'traffic safety act'],
            'nova scotia': ['nova scotia', 'motor vehicle act'],
            'new brunswick': ['new brunswick', 'motor vehicle act'],
            'newfoundland': ['newfoundland', 'highway traffic act'],
            'prince edward island': ['prince edward island', 'highway traffic act'],
            'northwest territories': ['northwest territories'],
            'yukon': ['yukon'],
            'nunavut': ['nunavut']
        }

        # US states (simplified)
        us_states = [
            'california', 'texas', 'florida', 'new york', 'pennsylvania',
            'illinois', 'ohio', 'georgia', 'north carolina', 'michigan'
        ]

        text_lower = text.lower()

        # Check Canadian provinces
        for province, keywords in provinces.items():
            for keyword in keywords:
                if keyword in text_lower:
                    return province.title()

        # Check US states
        for state in us_states:
            if state in text_lower:
                return state.title()

        return None

    def chunk_text(
        self,
        text: str,
        metadata: Optional[Dict[str, Any]] = None
    ) -> List[Dict[str, Any]]:
        """
        Chunk text into smaller pieces with metadata.

        Args:
            text: Text content to chunk
            metadata: Base metadata to include in each chunk

        Returns:
            List of chunk dictionaries with content and metadata
        """
        if not text or not text.strip():
            return []

        # Split text into chunks
        chunks = self.text_splitter.split_text(text)

        chunk_list = []
        base_metadata = metadata or {}

        for i, chunk_text in enumerate(chunks):
            # Skip empty chunks
            if not chunk_text.strip():
                continue

            chunk_metadata = base_metadata.copy()
            chunk_metadata.update({
                'chunk_index': i,
                'chunk_id': f"{base_metadata.get('doc_id', 'doc')}_chunk_{i}",
                'content_length': len(chunk_text)
            })

            chunk_list.append({
                'content': chunk_text,
                'metadata': chunk_metadata
            })

        logger.debug(f"📄 Created {len(chunk_list)} chunks from {len(text)} characters")
        return chunk_list

    def process_document(
        self,
        file_path: str,
        doc_id: Optional[str] = None,
        user_id: Optional[str] = None
    ) -> Dict[str, Any]:
        """
        Process a document and return structured data.

        Args:
            file_path: Path to document file
            doc_id: Document ID (auto-generated if None)
            user_id: User ID for metadata

        Returns:
            Dictionary with extracted content, chunks, and metadata
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"Document not found: {file_path}")

        file_path_obj = Path(file_path)
        file_name = file_path_obj.name
        file_ext = file_path_obj.suffix.lower()

        # Generate doc_id if not provided
        if not doc_id:
            import uuid
            doc_id = f"doc_{uuid.uuid4().hex[:16]}"

        logger.info(f"📄 Processing document: {file_name}")

        # Initialize result structure
        result = {
            'doc_id': doc_id,
            'file_name': file_name,
            'file_path': str(file_path),
            'file_type': file_ext,
            'file_size': os.path.getsize(file_path),
            'text_chunks': [],
            'tables': [],
            'detected_offence_number': None,
            'detected_province': None,
            'total_chunks': 0,
            'processing_metadata': {}
        }

        # Process based on file type
        if file_ext == '.pdf':
            text_content = self.extract_text_from_pdf(file_path)
            result['processing_metadata']['extraction_method'] = 'pdfplumber'

        elif file_ext == '.docx':
            text_content = self.extract_text_from_docx(file_path)
            result['processing_metadata']['extraction_method'] = 'python-docx'

        elif file_ext in ['.xlsx', '.xls']:
            text_content, tables = self.extract_text_from_xlsx(file_path)
            result['tables'] = tables
            result['processing_metadata']['extraction_method'] = 'openpyxl'

        elif file_ext in ['.jpg', '.jpeg', '.png', '.bmp', '.tiff']:
            # Try OCR
            text_content = self.extract_text_from_image(file_path)
            if not text_content and EASYOCR_AVAILABLE:
                text_content = self.extract_text_from_image_easyocr(file_path)
            result['processing_metadata']['extraction_method'] = 'ocr'

        elif file_ext == '.txt':
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                text_content = f.read()
            result['processing_metadata']['extraction_method'] = 'plain_text'

        else:
            # Try to read as text
            try:
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    text_content = f.read()
                result['processing_metadata']['extraction_method'] = 'fallback_text'
            except:
                raise ValueError(f"Unsupported file type: {file_ext}")

        # Detect offence number and province
        if text_content:
            result['detected_offence_number'] = self.detect_offence_number(text_content)
            result['detected_province'] = self.detect_province(text_content)

        # Create base metadata for chunks
        base_metadata = {
            'doc_id': doc_id,
            'file_name': file_name,
            'file_type': file_ext,
            'user_id': user_id,
            'province': result['detected_province'],
            'offence_number': result['detected_offence_number']
        }

        # Chunk the text
        if text_content:
            result['text_chunks'] = self.chunk_text(text_content, base_metadata)
            result['total_chunks'] = len(result['text_chunks'])

        logger.info(f"✅ Document processed: {result['total_chunks']} chunks, "
                   f"offence: {result['detected_offence_number']}, "
                   f"province: {result['detected_province']}")

        return result

    def process_batch(
        self,
        file_paths: List[str],
        user_id: Optional[str] = None
    ) -> List[Dict[str, Any]]:
        """
        Process multiple documents in batch.

        Args:
            file_paths: List of file paths to process
            user_id: User ID for metadata

        Returns:
            List of processed document results
        """
        results = []
        for file_path in file_paths:
            try:
                result = self.process_document(file_path, user_id=user_id)
                results.append(result)
            except Exception as e:
                logger.error(f"Failed to process {file_path}: {e}")
                results.append({
                    'doc_id': f"error_{Path(file_path).name}",
                    'file_name': Path(file_path).name,
                    'error': str(e),
                    'total_chunks': 0
                })

        return results

    def get_processor_info(self) -> Dict[str, Any]:
        """Get information about available processing capabilities."""
        return {
            'chunk_size': self.chunk_size,
            'chunk_overlap': self.chunk_overlap,
            'separators': self.separators,
            'capabilities': {
                'pdf': PDFPLUMBER_AVAILABLE or pypdf2 is not None,
                'docx': DOCX_AVAILABLE,
                'xlsx': OPENPYXL_AVAILABLE,
                'ocr': OCR_AVAILABLE,
                'easyocr': EASYOCR_AVAILABLE
            },
            'offence_patterns': len(self.offence_patterns)
        }


# Global processor instance
_processor_instance = None

def get_document_processor(
    chunk_size: int = 1000,
    chunk_overlap: int = 200
) -> DocumentProcessor:
    """Get or create global document processor instance."""
    global _processor_instance
    if _processor_instance is None:
        _processor_instance = DocumentProcessor(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap
        )
    return _processor_instance