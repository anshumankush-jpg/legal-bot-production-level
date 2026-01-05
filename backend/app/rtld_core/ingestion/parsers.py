"""
Document parsers for different file types - extracted from artillty
"""

import re
import uuid
from typing import List, Dict, Optional, Any, Tuple
from pathlib import Path
import logging

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import openpyxl
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False

try:
    import pandas as pd
    PANDAS_AVAILABLE = True
except ImportError:
    PANDAS_AVAILABLE = False

from ..schemas import Chunk

logger = logging.getLogger(__name__)


class RTLDTextParser:
    """Parse text documents into chunks"""

    def __init__(self, chunk_size: int = 1000, overlap: int = 200):
        self.chunk_size = chunk_size
        self.overlap = overlap

    def parse_text(self, text: str, doc_id: str, metadata: Optional[Dict[str, Any]] = None) -> List[Chunk]:
        """Parse text into chunks"""
        if not text.strip():
            return []

        metadata = metadata or {}

        # Split into paragraphs first
        paragraphs = [p.strip() for p in re.split(r'\n\s*\n', text) if p.strip()]

        if not paragraphs:
            return []

        # If content is small, return as single chunk
        if len(text) <= self.chunk_size:
            chunk = Chunk(
                id=f"{doc_id}_chunk_0",
                content=text,
                metadata={**metadata, 'doc_id': doc_id, 'chunk_index': 0},
                doc_id=doc_id,
                chunk_index=0
            )
            return [chunk]

        # Chunk larger documents
        chunks = []
        chunk_texts = self._chunk_text(text)

        for i, chunk_text in enumerate(chunk_texts):
            chunk = Chunk(
                id=f"{doc_id}_chunk_{i}",
                content=chunk_text,
                metadata={**metadata, 'doc_id': doc_id, 'chunk_index': i},
                doc_id=doc_id,
                chunk_index=i
            )
            chunks.append(chunk)

        return chunks

    def _chunk_text(self, text: str) -> List[str]:
        """Split text into overlapping chunks"""
        if len(text) <= self.chunk_size:
            return [text]

        chunks = []
        start = 0

        while start < len(text):
            end = start + self.chunk_size

            # If we're not at the end, try to find a good breaking point
            if end < len(text):
                # Look for sentence endings within the last 100 characters
                last_period = text.rfind('.', start, end)
                last_newline = text.rfind('\n', start, end)

                # Use the latest good breaking point
                break_point = max(last_period, last_newline)
                if break_point > start + self.chunk_size - 100:
                    end = break_point + 1
                else:
                    # Force break at chunk_size
                    pass

            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)

            # Move start forward
            start = end - self.overlap

            # Prevent infinite loop
            if start >= len(text):
                break

        return chunks


class RTLDDocxParser:
    """Parse DOCX documents"""

    def __init__(self, chunk_size: int = 1000, overlap: int = 200):
        self.text_parser = RTLDTextParser(chunk_size, overlap)

    def parse_docx(self, file_path: str, doc_id: str, metadata: Optional[Dict[str, Any]] = None) -> List[Chunk]:
        """Parse DOCX file into chunks"""
        if not DOCX_AVAILABLE:
            logger.warning("python-docx not available, cannot parse DOCX")
            return []

        try:
            doc = DocxDocument(file_path)
            text = ""

            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n\n"

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        text += row_text + "\n"

            if not text.strip():
                return []

            return self.text_parser.parse_text(text, doc_id, metadata)

        except Exception as e:
            logger.error(f"Error parsing DOCX {file_path}: {e}")
            return []


class RTLDExcelParser:
    """Parse Excel documents"""

    def __init__(self, chunk_size: int = 1000, overlap: int = 200):
        self.text_parser = RTLDTextParser(chunk_size, overlap)

    def parse_excel(self, file_path: str, doc_id: str, metadata: Optional[Dict[str, Any]] = None) -> List[Chunk]:
        """Parse Excel file into chunks"""
        if not EXCEL_AVAILABLE or not PANDAS_AVAILABLE:
            logger.warning("openpyxl/pandas not available, cannot parse Excel")
            return []

        try:
            # Read all sheets
            all_text = ""

            # Try to read as CSV first (simpler)
            if Path(file_path).suffix.lower() == '.csv':
                df = pd.read_csv(file_path)
            else:
                df = pd.read_excel(file_path)

            # Convert each row to text representation
            for _, row in df.iterrows():
                row_values = [str(val) for val in row.values if pd.notna(val)]
                if row_values:
                    row_text = " | ".join(row_values)
                    all_text += row_text + "\n"

            if not all_text.strip():
                return []

            return self.text_parser.parse_text(all_text, doc_id, metadata)

        except Exception as e:
            logger.error(f"Error parsing Excel {file_path}: {e}")
            return []


class RTLDUnifiedParser:
    """Unified parser that handles multiple document types"""

    def __init__(self, chunk_size: int = 1000, overlap: int = 200):
        self.chunk_size = chunk_size
        self.overlap = overlap
        self.text_parser = RTLDTextParser(chunk_size, overlap)
        self.docx_parser = RTLDDocxParser(chunk_size, overlap)
        self.excel_parser = RTLDExcelParser(chunk_size, overlap)

    def parse_file(self, file_path: str, doc_id: str, metadata: Optional[Dict[str, Any]] = None) -> Tuple[List[Chunk], str]:
        """
        Parse file and return chunks + content type

        Returns:
            Tuple of (chunks, content_type)
        """
        file_path = Path(file_path)
        file_ext = file_path.suffix.lower()
        metadata = metadata or {}

        # Add file info to metadata
        metadata.update({
            'filename': file_path.name,
            'file_path': str(file_path),
            'file_size': file_path.stat().st_size if file_path.exists() else 0
        })

        if file_ext in ['.txt', '.md']:
            # Read and parse text file
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    text = f.read()
                chunks = self.text_parser.parse_text(text, doc_id, metadata)
                return chunks, 'text'
            except Exception as e:
                logger.error(f"Error reading text file {file_path}: {e}")
                return [], 'text'

        elif file_ext in ['.docx', '.doc']:
            chunks = self.docx_parser.parse_docx(str(file_path), doc_id, metadata)
            return chunks, 'document'

        elif file_ext in ['.xlsx', '.xls', '.csv']:
            chunks = self.excel_parser.parse_excel(str(file_path), doc_id, metadata)
            return chunks, 'table'

        elif file_ext in ['.pdf']:
            # PDF parsing is handled by OCR processor
            # Return empty chunks for now (will be handled by OCR)
            return [], 'document'

        elif file_ext in ['.jpg', '.jpeg', '.png', '.bmp', '.tiff']:
            # Image parsing is handled by OCR processor
            return [], 'image'

        else:
            logger.warning(f"Unsupported file type: {file_ext}")
            return [], 'unknown'


class DocumentParser:
    """Generic document parser interface."""

    def __init__(self, chunk_size: int = 1000, overlap: int = 200):
        self.unified_parser = RTLDUnifiedParser(chunk_size, overlap)

    def parse_document(
        self,
        file_bytes: bytes,
        filename: str,
        doc_id: str,
        metadata: Optional[Dict[str, Any]] = None
    ) -> Tuple[List[Chunk], str]:
        """Parse document from bytes and filename."""
        # Save to temp file and parse
        import tempfile
        import os

        with tempfile.NamedTemporaryFile(delete=False, suffix=Path(filename).suffix) as tmp:
            tmp.write(file_bytes)
            tmp_path = tmp.name

        try:
            chunks, content_type = self.unified_parser.parse_file(tmp_path, doc_id, metadata)
            return chunks, content_type
        finally:
            try:
                os.unlink(tmp_path)
            except:
                pass


def parse_document(
    file_bytes: bytes,
    filename: str,
    doc_id: str,
    chunk_size: int = 1000,
    overlap: int = 200,
    metadata: Optional[Dict[str, Any]] = None
) -> Tuple[List[Chunk], str]:
    """Convenience function to parse a document from bytes."""
    parser = DocumentParser(chunk_size, overlap)
    return parser.parse_document(file_bytes, filename, doc_id, metadata)