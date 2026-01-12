"""
Unified Multi-Modal Embedding Server
Handles Text, Images, Tables, and Documents with automatic content extraction
"""

import os
import io
import tempfile
from typing import List, Dict, Optional, Union, Any
from pathlib import Path
import numpy as np
import faiss
from PIL import Image
try:
    import clip
    CLIP_AVAILABLE = True
except ImportError:
    CLIP_AVAILABLE = False
    import sys
    if sys.stdout.encoding and 'utf' in sys.stdout.encoding.lower():
        print("⚠️ CLIP not available. Install with: pip install git+https://github.com/openai/CLIP.git")
    else:
        print("WARNING: CLIP not available. Install with: pip install git+https://github.com/openai/CLIP.git")
import torch
from sentence_transformers import SentenceTransformer
import pandas as pd
from pydantic import BaseModel

# Document processing imports
try:
    import PyPDF2 as pypdf2
except ImportError:
    pypdf2 = None
import pdfplumber
from docx import Document as DocxDocument
import openpyxl


class EmbeddingRequest(BaseModel):
    """Request model for embedding API"""
    content: Optional[str] = None
    file_path: Optional[str] = None
    content_type: Optional[str] = None  # 'text', 'image', 'table', 'document', 'auto'
    metadata: Optional[Dict[str, Any]] = None


class EmbeddingResponse(BaseModel):
    """Response model for embedding API"""
    embeddings: List[List[float]]
    chunk_ids: List[str]
    chunk_texts: List[str]
    content_type: str
    num_chunks: int
    metadata: Optional[Dict[str, Any]] = None


class UnifiedEmbeddingServer:
    """
    Unified embedding server that handles:
    - Text: Using SentenceTransformer (best from tests)
    - Images: Using CLIP (multi-modal)
    - Tables: Converted to text and embedded
    - Documents: Auto-extracted and embedded (PDF, DOCX, PPTX, etc.)
    """
    
    def __init__(
        self,
        text_model_name: str = "all-MiniLM-L6-v2",
        image_model_name: str = "ViT-B/32",
        embedding_dim: int = 384,
        device: Optional[str] = None
    ):
        """
        Initialize the unified embedding server
        
        Args:
            text_model_name: SentenceTransformer model (default: winner from tests)
            image_model_name: CLIP model for images
            embedding_dim: Dimension of embeddings (384 for MiniLM, 512 for CLIP)
            device: 'cuda' or 'cpu'
        """
        self.device = device or ("cuda" if torch.cuda.is_available() else "cpu")
        import sys
        use_emoji = sys.stdout.encoding and 'utf' in sys.stdout.encoding.lower()
        if use_emoji:
            print(f"🚀 Initializing Unified Embedding Server on {self.device}")
        else:
            print(f"[*] Initializing Unified Embedding Server on {self.device}")
        
        # Initialize text embedding model (SentenceTransformer - winner from tests)
        if use_emoji:
            print("📝 Loading text embedding model (SentenceTransformer)...")
        else:
            print("[*] Loading text embedding model (SentenceTransformer)...")
        self.text_model = SentenceTransformer(text_model_name, device=self.device)
        self.text_embedding_dim = self.text_model.get_sentence_embedding_dimension()
        
        # Initialize image embedding model (CLIP)
        if CLIP_AVAILABLE:
            if use_emoji:
                print("🖼️ Loading image embedding model (CLIP)...")
            else:
                print("[*] Loading image embedding model (CLIP)...")
            self.clip_model, self.clip_preprocess = clip.load(image_model_name, device=self.device)
            self.image_embedding_dim = 512  # CLIP ViT-B/32 produces 512-dim embeddings
        else:
            if use_emoji:
                print("⚠️ CLIP not available - image embedding disabled")
            else:
                print("[!] CLIP not available - image embedding disabled")
            self.clip_model = None
            self.clip_preprocess = None
            self.image_embedding_dim = 512
        
        # Use text embedding dim as base (can align others)
        self.embedding_dim = self.text_embedding_dim
        
        # Initialize FAISS index (winner from tests)
        if use_emoji:
            print("💾 Initializing FAISS vector database...")
        else:
            print("[*] Initializing FAISS vector database...")
        self.index = faiss.IndexFlatIP(self.embedding_dim)  # Inner product for cosine similarity
        self.metadata_store: Dict[int, Dict] = {}  # Store metadata for each vector
        self.next_id = 0
        
        if use_emoji:
            print("✅ Unified Embedding Server initialized!")
        else:
            print("[+] Unified Embedding Server initialized!")
    
    def detect_content_type(self, content: Optional[str] = None, file_path: Optional[str] = None) -> str:
        """
        Auto-detect content type from content or file extension
        """
        if file_path:
            ext = Path(file_path).suffix.lower()
            if ext in ['.pdf', '.docx', '.pptx', '.txt', '.md']:
                return 'document'
            elif ext in ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.webp']:
                return 'image'
            elif ext in ['.csv', '.xlsx', '.xls', '.tsv']:
                return 'table'
            else:
                return 'text'
        
        if content:
            # Try to detect if it's structured data
            if '\t' in content or content.count(',') > 5:
                return 'table'
            # Check if it's image data (base64, etc.)
            if content.startswith('data:image') or content.startswith('/9j/'):
                return 'image'
            return 'text'
        
        return 'text'
    
    def embed_text(self, texts: List[str]) -> np.ndarray:
        """
        Embed text using SentenceTransformer (winner from tests)
        """
        if not texts:
            return np.array([])
        
        embeddings = self.text_model.encode(
            texts,
            convert_to_numpy=True,
            normalize_embeddings=True  # For cosine similarity
        )
        return embeddings
    
    def embed_image(self, image_path: str) -> np.ndarray:
        """
        Embed image using CLIP
        """
        if not CLIP_AVAILABLE or self.clip_model is None:
            raise ValueError("CLIP is not available. Install with: pip install git+https://github.com/openai/CLIP.git")
        
        image = Image.open(image_path).convert('RGB')
        image_tensor = self.clip_preprocess(image).unsqueeze(0).to(self.device)
        
        with torch.no_grad():
            image_features = self.clip_model.encode_image(image_tensor)
            image_features = image_features / image_features.norm(dim=-1, keepdim=True)  # Normalize
        
        # Project to text embedding space if needed
        embedding = image_features.cpu().numpy()
        
        # If dimensions don't match, use a simple projection
        if embedding.shape[1] != self.embedding_dim:
            # Use first N dimensions or pad/truncate
            if embedding.shape[1] > self.embedding_dim:
                embedding = embedding[:, :self.embedding_dim]
            else:
                padding = np.zeros((embedding.shape[0], self.embedding_dim - embedding.shape[1]))
                embedding = np.concatenate([embedding, padding], axis=1)
        
        return embedding
    
    def embed_table(self, table_data: Union[str, pd.DataFrame, str]) -> np.ndarray:
        """
        Embed table by converting to text description
        """
        if isinstance(table_data, str):
            # Try to parse as CSV
            try:
                df = pd.read_csv(io.StringIO(table_data))
            except:
                # If not CSV, treat as text
                return self.embed_text([table_data])
        elif isinstance(table_data, pd.DataFrame):
            df = table_data
        else:
            return self.embed_text([str(table_data)])
        
        # Convert table to descriptive text
        text_descriptions = []
        for idx, row in df.iterrows():
            # Create a text description of the row
            row_text = f"Row {idx}: " + ", ".join([f"{col}={val}" for col, val in row.items()])
            text_descriptions.append(row_text)
        
        # Also create a summary
        summary = f"Table with {len(df)} rows and {len(df.columns)} columns: {', '.join(df.columns)}"
        text_descriptions.insert(0, summary)
        
        return self.embed_text(text_descriptions)
    
    def extract_document_content(self, file_path: str) -> Dict[str, Any]:
        """
        Extract content from documents (PDF, DOCX, PPTX, etc.)
        Returns chunks of text, images, and tables
        """
        ext = Path(file_path).suffix.lower()
        chunks = []
        images = []
        tables = []
        
        if ext == '.pdf':
            # Extract text from PDF
            try:
                with pdfplumber.open(file_path) as pdf:
                    for page_num, page in enumerate(pdf.pages):
                        text = page.extract_text()
                        if text:
                            chunks.append({
                                'type': 'text',
                                'content': text,
                                'page': page_num + 1,
                                'chunk_id': f"pdf_page_{page_num + 1}"
                            })
                        
                        # Extract tables
                        page_tables = page.extract_tables()
                        for table_num, table in enumerate(page_tables):
                            if table:
                                df = pd.DataFrame(table[1:], columns=table[0])
                                tables.append({
                                    'type': 'table',
                                    'content': df,
                                    'page': page_num + 1,
                                    'table_num': table_num + 1,
                                    'chunk_id': f"pdf_page_{page_num + 1}_table_{table_num + 1}"
                                })
            except Exception as e:
                print(f"[!] Error extracting PDF: {e}")
                # Fallback to PyPDF2
                if pypdf2:
                    try:
                        with open(file_path, 'rb') as f:
                            pdf_reader = pypdf2.PdfReader(f)
                            for page_num, page in enumerate(pdf_reader.pages):
                                text = page.extract_text()
                                if text:
                                    chunks.append({
                                        'type': 'text',
                                        'content': text,
                                        'page': page_num + 1,
                                        'chunk_id': f"pdf_page_{page_num + 1}"
                                    })
                    except Exception as e2:
                        print(f"[!] Fallback PDF extraction also failed: {e2}")
                else:
                    print("[!] PyPDF2 not available for fallback")
        
        elif ext == '.docx':
            # Extract from DOCX
            try:
                doc = DocxDocument(file_path)
                full_text = []
                for para_num, para in enumerate(doc.paragraphs):
                    if para.text.strip():
                        full_text.append(para.text)
                        chunks.append({
                            'type': 'text',
                            'content': para.text,
                            'paragraph': para_num + 1,
                            'chunk_id': f"docx_para_{para_num + 1}"
                        })
                
                # Extract tables from DOCX
                for table_num, table in enumerate(doc.tables):
                    table_data = []
                    for row in table.rows:
                        table_data.append([cell.text for cell in row.cells])
                    if table_data:
                        df = pd.DataFrame(table_data[1:], columns=table_data[0])
                        tables.append({
                            'type': 'table',
                            'content': df,
                            'table_num': table_num + 1,
                            'chunk_id': f"docx_table_{table_num + 1}"
                        })
            except Exception as e:
                print(f"[!] Error extracting DOCX: {e}")
        
        elif ext in ['.xlsx', '.xls']:
            # Extract from Excel
            try:
                excel_file = pd.ExcelFile(file_path)
                for sheet_name in excel_file.sheet_names:
                    df = pd.read_excel(file_path, sheet_name=sheet_name)
                    tables.append({
                        'type': 'table',
                        'content': df,
                        'sheet': sheet_name,
                        'chunk_id': f"excel_sheet_{sheet_name}"
                    })
            except Exception as e:
                print(f"[!] Error extracting Excel: {e}")
        
        elif ext == '.csv':
            # Extract from CSV
            try:
                df = pd.read_csv(file_path)
                tables.append({
                    'type': 'table',
                    'content': df,
                    'chunk_id': "csv_main"
                })
            except Exception as e:
                print(f"[!] Error extracting CSV: {e}")
        
        elif ext in ['.txt', '.md']:
            # Plain text
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()
                    # Split into chunks (by paragraphs or fixed size)
                    paragraphs = content.split('\n\n')
                    for para_num, para in enumerate(paragraphs):
                        if para.strip():
                            chunks.append({
                                'type': 'text',
                                'content': para,
                                'chunk_id': f"text_para_{para_num + 1}"
                            })
            except Exception as e:
                print(f"[!] Error extracting text: {e}")
        
        return {
            'text_chunks': chunks,
            'tables': tables,
            'images': images  # TODO: Extract images from documents
        }
    
    def embed_document(self, file_path: str) -> EmbeddingResponse:
        """
        Embed a document by extracting content and embedding each chunk
        """
        # print(f"📄 Processing document: {file_path}")  # Commented to avoid encoding issues
        extracted = self.extract_document_content(file_path)
        
        all_embeddings = []
        all_chunk_ids = []
        all_chunk_texts = []
        
        # Embed text chunks
        if extracted['text_chunks']:
            texts = [chunk['content'] for chunk in extracted['text_chunks']]
            text_embeddings = self.embed_text(texts)
            all_embeddings.append(text_embeddings)
            all_chunk_ids.extend([chunk['chunk_id'] for chunk in extracted['text_chunks']])
            all_chunk_texts.extend(texts)
        
        # Embed tables
        if extracted['tables']:
            for table_info in extracted['tables']:
                table_emb = self.embed_table(table_info['content'])
                all_embeddings.append(table_emb)
                all_chunk_ids.append(table_info['chunk_id'])
                all_chunk_texts.append(f"Table: {table_info.get('sheet', '')} with {len(table_info['content'])} rows")
        
        if not all_embeddings:
            raise ValueError(f"No content extracted from document: {file_path}")
        
        # Concatenate all embeddings
        final_embeddings = np.concatenate(all_embeddings, axis=0) if len(all_embeddings) > 1 else all_embeddings[0]
        
        return EmbeddingResponse(
            embeddings=final_embeddings.tolist(),
            chunk_ids=all_chunk_ids,
            chunk_texts=all_chunk_texts,
            content_type='document',
            num_chunks=len(all_chunk_ids),
            metadata={'file_path': file_path, 'extracted': len(extracted['text_chunks']), 'tables': len(extracted['tables'])}
        )
    
    def embed(self, request: EmbeddingRequest) -> EmbeddingResponse:
        """
        Main embedding method - auto-detects content type and routes appropriately
        """
        # Auto-detect content type if not specified
        content_type = request.content_type or self.detect_content_type(
            request.content, request.file_path
        )
        
        if content_type == 'document' and request.file_path:
            return self.embed_document(request.file_path)
        
        elif content_type == 'image' and request.file_path:
            embedding = self.embed_image(request.file_path)
            return EmbeddingResponse(
                embeddings=[embedding[0].tolist()],
                chunk_ids=["image_0"],
                chunk_texts=[f"Image from {request.file_path}"],
                content_type='image',
                num_chunks=1,
                metadata=request.metadata
            )
        
        elif content_type == 'table' and request.file_path:
            df = pd.read_csv(request.file_path) if Path(request.file_path).suffix == '.csv' else pd.read_excel(request.file_path)
            embedding = self.embed_table(df)
            return EmbeddingResponse(
                embeddings=embedding.tolist(),
                chunk_ids=["table_0"],
                chunk_texts=[f"Table with {len(df)} rows"],
                content_type='table',
                num_chunks=1,
                metadata=request.metadata
            )
        
        elif content_type == 'text' and request.content:
            # Handle text - can be single or multiple chunks
            texts = [request.content] if isinstance(request.content, str) else request.content
            embeddings = self.embed_text(texts)
            chunk_ids = [f"text_{i}" for i in range(len(texts))]
            
            return EmbeddingResponse(
                embeddings=embeddings.tolist(),
                chunk_ids=chunk_ids,
                chunk_texts=texts,
                content_type='text',
                num_chunks=len(texts),
                metadata=request.metadata
            )
        
        else:
            raise ValueError(f"Cannot process content type: {content_type} with provided inputs")
    
    def add_to_index(self, embeddings: np.ndarray, metadata: List[Dict]) -> List[int]:
        """
        Add embeddings to FAISS index
        """
        if embeddings.shape[0] == 0:
            return []
        
        # Ensure float32 and contiguous
        embeddings = np.ascontiguousarray(embeddings.astype('float32'))
        
        # Normalize for cosine similarity
        faiss.normalize_L2(embeddings)
        
        # Add to index
        start_id = self.next_id
        self.index.add(embeddings)
        
        # Store metadata
        ids = []
        for i, meta in enumerate(metadata):
            vec_id = start_id + i
            self.metadata_store[vec_id] = meta
            ids.append(vec_id)
        
        self.next_id += len(embeddings)
        return ids
    
    def search(self, query: str, k: int = 5) -> List[Dict]:
        """
        Search for similar content
        """
        # Embed query
        query_emb = self.embed_text([query])
        query_emb = np.ascontiguousarray(query_emb.astype('float32'))
        faiss.normalize_L2(query_emb)
        
        # Search
        distances, indices = self.index.search(query_emb, k)
        
        results = []
        for dist, idx in zip(distances[0], indices[0]):
            if idx >= 0 and idx in self.metadata_store:
                results.append({
                    'id': int(idx),
                    'similarity': float(dist),
                    'metadata': self.metadata_store[idx]
                })
        
        return results
    
    def save_index(self, path: str):
        """Save FAISS index to disk"""
        faiss.write_index(self.index, path)
        # Save metadata separately
        import json
        metadata_path = path.replace('.index', '_metadata.json')
        with open(metadata_path, 'w') as f:
            json.dump(self.metadata_store, f, default=str)
    
    def load_index(self, path: str):
        """Load FAISS index from disk"""
        self.index = faiss.read_index(path)
        # Load metadata
        import json
        metadata_path = path.replace('.index', '_metadata.json')
        if os.path.exists(metadata_path):
            with open(metadata_path, 'r') as f:
                self.metadata_store = json.load(f)
                # Convert keys back to int
                self.metadata_store = {int(k): v for k, v in self.metadata_store.items()}
            self.next_id = max(self.metadata_store.keys()) + 1 if self.metadata_store else 0


if __name__ == "__main__":
    # Example usage
    server = UnifiedEmbeddingServer()
    
    # Test text embedding
    print("\n📝 Testing text embedding...")
    text_request = EmbeddingRequest(content="This is a test document about artificial intelligence.")
    text_response = server.embed(text_request)
    print(f"✅ Text embedded: {len(text_response.embeddings)} chunks, dim={len(text_response.embeddings[0])}")
    
    # Add to index
    embeddings = np.array(text_response.embeddings)
    metadata = [{'chunk_id': cid, 'text': text, 'type': 'text'} 
                for cid, text in zip(text_response.chunk_ids, text_response.chunk_texts)]
    server.add_to_index(embeddings, metadata)
    
    # Test search
    print("\n🔍 Testing search...")
    results = server.search("machine learning", k=1)
    print(f"✅ Search results: {results}")

