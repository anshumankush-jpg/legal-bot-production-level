"""
Artillery Document Processor for PLAZA-AI
Handles PDF, DOCX, TXT files with text extraction, OCR, chunking, and offence number detection
"""

import os
import re
import io
import tempfile
from pathlib import Path
from typing import List, Dict, Optional, Any, Tuple
import logging

# Document processing libraries
try:
    import PyPDF2 as pypdf2
except ImportError:
    pypdf2 = None

try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False

try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import openpyxl
    OPENPYXL_AVAILABLE = True
except ImportError:
    OPENPYXL_AVAILABLE = False

try:
    import pytesseract
    from PIL import Image
    OCR_AVAILABLE = True
    
    # Configure Tesseract path explicitly for Windows
    if os.name == 'nt':  # Windows
        tesseract_path = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
        if os.path.exists(tesseract_path):
            pytesseract.pytesseract.tesseract_cmd = tesseract_path
            logging.info(f"[OCR] Tesseract configured at: {tesseract_path}")
            # Test if Tesseract actually works
            try:
                version = pytesseract.get_tesseract_version()
                logging.info(f"[OCR] Tesseract version: {version}")
            except Exception as e:
                logging.warning(f"[OCR] Tesseract configured but not working: {e}")
                OCR_AVAILABLE = False
        else:
            logging.warning(f"[OCR] Tesseract not found at default path: {tesseract_path}")
            OCR_AVAILABLE = False
except ImportError as e:
    logging.warning(f"[OCR] Failed to import pytesseract: {e}")
    OCR_AVAILABLE = False

# Log OCR status at module load time
logging.info(f"[OCR] OCR_AVAILABLE = {OCR_AVAILABLE}")

try:
    import easyocr
    EASYOCR_AVAILABLE = True
except ImportError:
    EASYOCR_AVAILABLE = False

try:
    import pandas as pd
    PANDAS_AVAILABLE = True
except ImportError:
    PANDAS_AVAILABLE = False

logger = logging.getLogger(__name__)


class SimpleTextSplitter:
    """Simple text splitter to replace langchain dependency."""

    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.separators = ["\n\n", "\n", ". ", " ", ""]

    def split_text(self, text: str) -> List[str]:
        """Split text into chunks with overlap."""
        if not text:
            return []

        chunks = []
        start = 0

        while start < len(text):
            # Find the end of this chunk
            end = start + self.chunk_size

            # If we're not at the end, try to find a good break point
            if end < len(text):
                # Look for the best separator
                best_break = end
                for separator in self.separators:
                    if not separator:
                        continue
                    last_sep = text.rfind(separator, start, end)
                    if last_sep != -1 and last_sep > start:
                        best_break = last_sep + len(separator)
                        break

                end = best_break

            # Extract the chunk
            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)

            # Move start position with overlap
            start = end - self.chunk_overlap
            if start >= len(text):
                break

        return chunks


class ArtilleryDocumentProcessor:
    """
    Artillery Document Processor for PLAZA-AI.

    Features:
    - Multi-format document processing (PDF, DOCX, XLSX, images, text)
    - OCR for scanned documents
    - Intelligent text chunking with overlap
    - Offence number detection and province recognition
    - Table extraction and processing
    """

    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        """
        Initialize document processor.

        Args:
            chunk_size: Size of text chunks in characters
            chunk_overlap: Overlap between chunks in characters
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.separators = ["\n\n", "\n", ". ", " ", ""]

        # Initialize text splitter
        self.text_splitter = SimpleTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap
        )

        # Offence number patterns
        self.offence_patterns = [
            r'[Oo]ffence\s*[Nn]o\.?\s*:?\s*(\d{8,12})',
            r'[Oo]ffence\s*[Nn]umber\s*:?\s*(\d{8,12})',
            r'[Tt]icket\s*#?\s*:?\s*(\d{8,12})',
            r'[Cc]itation\s*#?\s*:?\s*(\d{8,12})',
            r'[Vv]iolation\s*#?\s*:?\s*(\d{8,12})',
            r'[Ii]nfraction\s*#?\s*:?\s*(\d{8,12})',
        ]

        # Province/state patterns
        self.province_patterns = {
            'ontario': ['ontario', 'highway traffic act', 'hta'],
            'quebec': ['quebec', 'code de la sécurité routière'],
            'british columbia': ['british columbia', 'motor vehicle act'],
            'alberta': ['alberta', 'traffic safety act'],
            'manitoba': ['manitoba', 'highway traffic act'],
            'saskatchewan': ['saskatchewan', 'traffic safety act'],
            'nova scotia': ['nova scotia', 'motor vehicle act'],
            'new brunswick': ['new brunswick', 'motor vehicle act'],
            'newfoundland': ['newfoundland', 'highway traffic act'],
            'prince edward island': ['prince edward island', 'highway traffic act'],
            'northwest territories': ['northwest territories'],
            'yukon': ['yukon'],
            'nunavut': ['nunavut']
        }

        logger.info("📄 Artillery Document Processor initialized")

    def extract_offence_number(self, text: str) -> Optional[str]:
        """
        Extract offence number from text using regex patterns.

        Args:
            text: Text content to search

        Returns:
            Detected offence number or None
        """
        for pattern in self.offence_patterns:
            matches = re.finditer(pattern, text, re.IGNORECASE)
            for match in matches:
                offence_num = match.group(1).strip()

                # Validate: should be 8-12 digits
                if len(offence_num) >= 8 and len(offence_num) <= 12 and offence_num.isdigit():
                    logger.info(f"✅ Detected offence number: {offence_num}")
                    return offence_num

        return None

    def detect_province(self, text: str) -> Optional[str]:
        """
        Detect Canadian province or US state from text.

        Args:
            text: Text content to search

        Returns:
            Detected province/state or None
        """
        text_lower = text.lower()

        # Check Canadian provinces
        for province, keywords in self.province_patterns.items():
            for keyword in keywords:
                if keyword in text_lower:
                    return province.title()

        # Check US states (simplified)
        us_states = [
            'california', 'texas', 'florida', 'new york', 'pennsylvania',
            'illinois', 'ohio', 'georgia', 'north carolina', 'michigan'
        ]

        for state in us_states:
            if state in text_lower:
                return state.title()

        return None

    def process_pdf(self, file_path: str) -> Dict[str, Any]:
        """
        Extract text, tables, and images from PDF.

        Args:
            file_path: Path to PDF file

        Returns:
            Dictionary with extracted content
        """
        chunks = []
        tables = []
        images = []

        try:
            # Extract text and tables using pdfplumber
            if PDFPLUMBER_AVAILABLE:
                with pdfplumber.open(file_path) as pdf:
                    for page_num, page in enumerate(pdf.pages):
                        # Extract text
                        text = page.extract_text()
                        if text and text.strip():
                            chunks.append({
                                'type': 'text',
                                'content': text.strip(),
                                'page': page_num + 1
                            })

                        # Extract tables
                        page_tables = page.extract_tables()
                        for table_idx, table in enumerate(page_tables):
                            if table and len(table) > 1:  # Skip empty tables
                                if PANDAS_AVAILABLE:
                                    try:
                                        df = pd.DataFrame(table[1:], columns=table[0])
                                        tables.append({
                                            'type': 'table',
                                            'content': df,
                                            'page': page_num + 1,
                                            'table_index': table_idx
                                        })
                                    except Exception as e:
                                        logger.warning(f"Failed to process table {table_idx} on page {page_num + 1}: {e}")
                                else:
                                    tables.append({
                                        'type': 'table',
                                        'content': table,
                                        'page': page_num + 1,
                                        'table_index': table_idx
                                    })

            # Extract images using PyMuPDF
            if PYMUPDF_AVAILABLE:
                doc = fitz.open(file_path)
                for page_num in range(len(doc)):
                    page = doc[page_num]
                    image_list = page.get_images(full=False)

                    for img_index, img in enumerate(image_list):
                        try:
                            xref = img[0]
                            base_image = doc.extract_image(xref)

                            if base_image and base_image['image']:
                                images.append({
                                    'type': 'image',
                                    'data': base_image['image'],
                                    'ext': base_image['ext'],
                                    'page': page_num + 1,
                                    'image_index': img_index
                                })
                        except Exception as e:
                            logger.warning(f"Failed to extract image {img_index} on page {page_num + 1}: {e}")

                doc.close()

        except Exception as e:
            logger.error(f"Failed to process PDF {file_path}: {e}")
            raise

        return {
            'text_chunks': chunks,
            'tables': tables,
            'images': images
        }

    def process_docx(self, file_path: str) -> Dict[str, Any]:
        """
        Extract text and tables from DOCX.

        Args:
            file_path: Path to DOCX file

        Returns:
            Dictionary with extracted content
        """
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx not available. Install with: pip install python-docx")

        chunks = []
        tables = []

        try:
            doc = DocxDocument(file_path)

            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    chunks.append({
                        'type': 'text',
                        'content': para.text.strip(),
                        'page': None  # DOCX doesn't have pages
                    })

            # Extract tables
            for table_idx, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        row_data.append(cell.text.strip())
                    table_data.append(row_data)

                if table_data:
                    if PANDAS_AVAILABLE:
                        try:
                            df = pd.DataFrame(table_data[1:], columns=table_data[0])
                            tables.append({
                                'type': 'table',
                                'content': df,
                                'page': None,
                                'table_index': table_idx
                            })
                        except Exception as e:
                            logger.warning(f"Failed to process DOCX table {table_idx}: {e}")
                            tables.append({
                                'type': 'table',
                                'content': table_data,
                                'page': None,
                                'table_index': table_idx
                            })
                    else:
                        tables.append({
                            'type': 'table',
                            'content': table_data,
                            'page': None,
                            'table_index': table_idx
                        })

        except Exception as e:
            logger.error(f"Failed to process DOCX {file_path}: {e}")
            raise

        return {
            'text_chunks': chunks,
            'tables': tables,
            'images': []  # DOCX images not implemented yet
        }

    def process_text(self, file_path: str) -> Dict[str, Any]:
        """
        Process plain text file.

        Args:
            file_path: Path to text file

        Returns:
            Dictionary with extracted content
        """
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()

            chunks = [{
                'type': 'text',
                'content': content,
                'page': None
            }]

        except Exception as e:
            logger.error(f"Failed to process text file {file_path}: {e}")
            chunks = []

        return {
            'text_chunks': chunks,
            'tables': [],
            'images': []
        }

    def process_xlsx(self, file_path: str) -> Dict[str, Any]:
        """
        Process Excel file.

        Args:
            file_path: Path to XLSX file

        Returns:
            Dictionary with extracted content
        """
        if not OPENPYXL_AVAILABLE:
            raise ImportError("openpyxl not available. Install with: pip install openpyxl")

        chunks = []
        tables = []

        try:
            workbook = openpyxl.load_workbook(file_path)

            for sheet_name in workbook.sheetnames:
                worksheet = workbook[sheet_name]

                # Extract table data
                table_data = []
                for row in worksheet.iter_rows(values_only=True):
                    if any(cell for cell in row if cell is not None):
                        table_data.append([str(cell) if cell is not None else "" for cell in row])

                if table_data:
                    if PANDAS_AVAILABLE:
                        try:
                            df = pd.DataFrame(table_data[1:], columns=table_data[0])
                            tables.append({
                                'type': 'table',
                                'content': df,
                                'sheet': sheet_name,
                                'page': None
                            })
                        except Exception as e:
                            logger.warning(f"Failed to process XLSX sheet {sheet_name}: {e}")
                            tables.append({
                                'type': 'table',
                                'content': table_data,
                                'sheet': sheet_name,
                                'page': None
                            })
                    else:
                        tables.append({
                            'type': 'table',
                            'content': table_data,
                            'sheet': sheet_name,
                            'page': None
                        })

                    # Convert table to text for chunking
                    sheet_text = f"Sheet: {sheet_name}\n"
                    for i, row in enumerate(table_data):
                        sheet_text += f"Row {i}: " + " | ".join(row) + "\n"

                    chunks.append({
                        'type': 'text',
                        'content': sheet_text,
                        'page': None
                    })

        except Exception as e:
            logger.error(f"Failed to process XLSX {file_path}: {e}")
            raise

        return {
            'text_chunks': chunks,
            'tables': tables,
            'images': []
        }

    def process_image(self, file_path: str) -> Dict[str, Any]:
        """
        Process image file with ENHANCED OCR.
        
        Now includes:
        - Image preprocessing (contrast, sharpening, denoising)
        - Automatic rotation correction
        - Pattern recognition (dates, codes, numbers)
        - Confidence scoring
        - Structured field extraction

        Args:
            file_path: Path to image file

        Returns:
            Dictionary with extracted content and metadata
        """
        chunks = []
        images = []
        metadata = {}

        try:
            # Try enhanced OCR first
            try:
                from artillery.enhanced_ocr import get_enhanced_ocr
                enhanced_ocr = get_enhanced_ocr()
                
                logger.info(f"[ARTILLERY] Using enhanced OCR for {Path(file_path).name}")
                ocr_result = enhanced_ocr.process_image_enhanced(file_path)
                
                if ocr_result.get('text'):
                    # Create main text chunk
                    content = ocr_result['text'].strip()
                    
                    # Add structured fields to the content for better searchability
                    if ocr_result.get('structured_fields'):
                        fields = ocr_result['structured_fields']
                        field_text = "\n\n[DETECTED FIELDS]\n"
                        if fields.get('dates'):
                            field_text += f"Dates: {', '.join(fields['dates'])}\n"
                        if fields.get('codes'):
                            field_text += f"Codes: {', '.join(fields['codes'])}\n"
                        if fields.get('numbers'):
                            field_text += f"Numbers: {', '.join(fields['numbers'][:10])}\n"
                        content += field_text
                    
                    # Add labeled fields
                    if ocr_result.get('labeled_fields'):
                        field_text = "\n[LABELED FIELDS]\n"
                        for label, value in ocr_result['labeled_fields'].items():
                            field_text += f"{label.replace('_', ' ').title()}: {value}\n"
                        content += field_text
                    
                    chunks.append({
                        'type': 'text',
                        'content': content,
                        'page': None,
                        'ocr_confidence': ocr_result.get('confidence', 0),
                        'preprocessing': ocr_result.get('preprocessing_used', 'unknown')
                    })
                    
                    # Store metadata
                    metadata['ocr_confidence'] = ocr_result.get('confidence', 0)
                    metadata['structured_fields'] = ocr_result.get('structured_fields', {})
                    metadata['labeled_fields'] = ocr_result.get('labeled_fields', {})
                    metadata['warnings'] = ocr_result.get('warnings', [])
                    metadata['suggestions'] = ocr_result.get('suggestions', [])
                    
                    logger.info(f"[ARTILLERY] Enhanced OCR: {len(content)} chars, {ocr_result.get('confidence', 0):.1f}% confidence")
                    
                else:
                    # Enhanced OCR failed, add error info
                    error_msg = f"[Image uploaded: {Path(file_path).name}]\n"
                    error_msg += f"OCR Status: {ocr_result.get('error', 'No text detected')}\n"
                    if ocr_result.get('warnings'):
                        error_msg += "Warnings:\n" + "\n".join(f"  • {w}" for w in ocr_result['warnings']) + "\n"
                    if ocr_result.get('suggestions'):
                        error_msg += "Suggestions:\n" + "\n".join(f"  • {s}" for s in ocr_result['suggestions'])
                    
                    chunks.append({
                        'type': 'text',
                        'content': error_msg,
                        'page': None
                    })
                    
            except ImportError:
                # Fall back to basic OCR
                logger.info(f"[ARTILLERY] Enhanced OCR not available, using basic OCR")
                if OCR_AVAILABLE:
                    image = Image.open(file_path)
                    text = pytesseract.image_to_string(image)

                    if text.strip():
                        chunks.append({
                            'type': 'text',
                            'content': text.strip(),
                            'page': None
                        })
                    else:
                        chunks.append({
                            'type': 'text',
                            'content': f"[Image uploaded: {Path(file_path).name}. No text detected.]",
                            'page': None
                        })
                else:
                    chunks.append({
                        'type': 'text',
                        'content': f"[Image uploaded: {Path(file_path).name}. OCR not available - install Tesseract to extract text]",
                        'page': None
                    })

            # Always include the image data for embedding
            with open(file_path, 'rb') as f:
                image_data = f.read()

            images.append({
                'type': 'image',
                'data': image_data,
                'ext': Path(file_path).suffix.lower().lstrip('.'),
                'page': None,
                'image_index': 0
            })

        except Exception as e:
            logger.error(f"Failed to process image {file_path}: {e}")
            # Don't raise - return what we have
            # Create minimal chunk so upload doesn't fail completely
            chunks.append({
                'type': 'text',
                'content': f"[Image uploaded: {Path(file_path).name}. Processing failed: {str(e)}]",
                'page': None
            })

        return {
            'text_chunks': chunks,
            'tables': [],
            'images': images
        }

    def chunk_document(self, content: str, metadata: Dict) -> List[Dict]:
        """
        Chunk text into smaller pieces with overlap.

        Args:
            content: Text content to chunk
            metadata: Base metadata to include in each chunk

        Returns:
            List of chunk dictionaries with content and metadata
        """
        if not content or not content.strip():
            return []

        # Split text into chunks
        chunks = self.text_splitter.split_text(content)

        chunk_list = []
        base_metadata = metadata.copy()

        for i, chunk_text in enumerate(chunks):
            # Skip empty chunks
            if not chunk_text.strip():
                continue

            chunk_metadata = base_metadata.copy()
            chunk_metadata.update({
                'chunk_index': i,
                'chunk_id': f"{base_metadata.get('doc_id', 'doc')}_chunk_{i}",
                'content_length': len(chunk_text)
            })

            chunk_list.append({
                'content': chunk_text,
                'metadata': chunk_metadata
            })

        logger.debug(f"📄 Created {len(chunk_list)} chunks from {len(content)} characters")
        return chunk_list

    def process_document(self, file_path: str) -> Dict[str, Any]:
        """
        Process a document based on its file extension.

        Args:
            file_path: Path to document file

        Returns:
            Dictionary with processed content
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"Document not found: {file_path}")

        file_path_obj = Path(file_path)
        file_ext = file_path_obj.suffix.lower()

        logger.info(f"📄 Processing document: {file_path_obj.name}")

        # Process based on file type
        if file_ext == '.pdf':
            result = self.process_pdf(file_path)
        elif file_ext == '.docx':
            result = self.process_docx(file_path)
        elif file_ext in ['.xlsx', '.xls']:
            result = self.process_xlsx(file_path)
        elif file_ext in ['.jpg', '.jpeg', '.png', '.bmp', '.tiff']:
            result = self.process_image(file_path)
        elif file_ext == '.txt':
            result = self.process_text(file_path)
        else:
            # Try to read as text
            try:
                result = self.process_text(file_path)
                logger.warning(f"Unknown file type {file_ext}, treating as text")
            except Exception:
                raise ValueError(f"Unsupported file type: {file_ext}")

        # Add file metadata
        result.update({
            'file_name': file_path_obj.name,
            'file_path': str(file_path),
            'file_type': file_ext,
            'file_size': os.path.getsize(file_path),
            'total_text_chunks': len(result['text_chunks']),
            'total_tables': len(result['tables']),
            'total_images': len(result['images'])
        })

        logger.info(f"✅ Document processed: {result['total_text_chunks']} text chunks, "
                   f"{result['total_tables']} tables, {result['total_images']} images")

        return result

    def get_processor_info(self) -> Dict[str, Any]:
        """Get information about available processing capabilities."""
        return {
            'chunk_size': self.chunk_size,
            'chunk_overlap': self.chunk_overlap,
            'separators': self.separators,
            'capabilities': {
                'pdf': PDFPLUMBER_AVAILABLE or pypdf2 is not None,
                'docx': DOCX_AVAILABLE,
                'xlsx': OPENPYXL_AVAILABLE,
                'ocr': OCR_AVAILABLE,
                'easyocr': EASYOCR_AVAILABLE,
                'tables': PANDAS_AVAILABLE
            },
            'offence_patterns': len(self.offence_patterns),
            'province_patterns': len(self.province_patterns)
        }


# Global processor instance
_processor_instance = None

def get_artillery_document_processor(
    chunk_size: int = 1000,
    chunk_overlap: int = 200
) -> ArtilleryDocumentProcessor:
    """Get or create global Artillery document processor instance."""
    global _processor_instance
    if _processor_instance is None:
        _processor_instance = ArtilleryDocumentProcessor(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap
        )
    return _processor_instance